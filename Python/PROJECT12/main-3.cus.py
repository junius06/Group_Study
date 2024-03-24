import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_run_style(run, text, font_name='나눔고딕', font_size=20, is_bold=False):
    run.text = text
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = docx.shared.Pt(font_size)
    run.bold = is_bold

def add_paragraph(doc, text='', font_name='나눔고딕', font_size=20, is_bold=False, alignment=None):
    para = doc.add_paragraph()
    run = para.add_run()
    set_run_style(run, text, font_name, font_size, is_bold)
    if alignment:
        para.alignment = alignment
    return para

doc = docx.Document(r'수료증양식.docx')

# Normal 스타일 변경
style = doc.styles['Normal']
style.font.name = '나눔고딕'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
style.font.size = docx.shared.Pt(12)

# 제목 추가
add_paragraph(doc, ' 제 1호\n ', font_size=40, alignment=WD_ALIGN_PARAGRAPH.CENTER)

# 항목 추가
add_paragraph(doc, '수 료 증', font_size=40, is_bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

# 정보 추가
info_texts = [
    ' 성 명 : 아 무 개 ',
    ' 생 년 월 일 : 2222.22.22 ',
    ' 교 육 과 정 : 파이썬과 40개의 작품들 ',
    ' 교 육 날 짜 : 1111.11.11 ~ 2222.22.22 '
]

for info_text in info_texts:
    add_paragraph(doc, info_text)

# 내용 추가
content_texts = [
    '위 사람은 파이썬과 40개의 작품들 교육 과정을 ',
    '이수하였으므로 이 증서를 수여한다. '
]

for content_text in content_texts:
    add_paragraph(doc, content_text)

# 날짜 추가
add_paragraph(doc, '2222.22.22', alignment=WD_ALIGN_PARAGRAPH.RIGHT)

# 교육기관 추가
add_paragraph(doc, '파이썬교육기관장', font_size=20, is_bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

doc.save(r'수료증결과.docx')